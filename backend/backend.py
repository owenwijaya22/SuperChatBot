import os
import sys
import uuid
import traceback
from typing import List

import boto3
import docx
import pymongo
from dotenv import load_dotenv
from PyPDF2 import PdfReader
from pydantic import BaseModel
from fastapi import FastAPI, UploadFile, status, HTTPException, Request
import time
from fastapi.middleware.cors import CORSMiddleware
from langchain.prompts import PromptTemplate
from langchain.chains import ConversationalRetrievalChain
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_openai import AzureChatOpenAI, AzureOpenAIEmbeddings
from langchain.docstore.document import Document as LangchainDocument
from langchain_community.callbacks.manager import get_openai_callback
from langchain_mongodb.vectorstores import MongoDBAtlasVectorSearch
from pymongo.operations import SearchIndexModel
from io import BytesIO


# from langchain.vectorstores.redis import Redis as RedisVectorStore

# redis_url = "redis://redis:6379"


if "OPENAI_API_BASE" in os.environ:
    del os.environ["OPENAI_API_BASE"]
load_dotenv()

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
AZURE_ENDPOINT = os.getenv("AZURE_ENDPOINT")
S3_KEY = os.getenv("S3_KEY")
S3_SECRET = os.getenv("S3_SECRET")
S3_BUCKET = os.getenv("S3_BUCKET")
S3_REGION = os.getenv("S3_REGION")
S3_PATH = os.getenv("S3_PATH")
MONGO_URL = os.getenv("MONGO_URL")

# Initialize Azure services
embeddings = AzureOpenAIEmbeddings(
    model="text-embedding-ada-002",
    azure_endpoint=AZURE_ENDPOINT,
    api_key=OPENAI_API_KEY,
    openai_api_version="2023-07-01-preview",
)

llm = AzureChatOpenAI(
    azure_endpoint=AZURE_ENDPOINT,
    openai_api_version="2023-07-01-preview",
    deployment_name="GPT4",
    openai_api_key=OPENAI_API_KEY,
    openai_api_type="azure",
    model_name="text-embedding-ada-002",
    temperature=0,
)

# Initialize AWS S3 to read file
s3 = boto3.client(
    "s3",
    aws_access_key_id=S3_KEY,
    aws_secret_access_key=S3_SECRET,
    region_name=S3_REGION,
)

# Initialize aws s3 session for uploading file
aws_s3 = boto3.Session(
    aws_access_key_id=S3_KEY,
    aws_secret_access_key=S3_SECRET,
    region_name=S3_REGION,
)

# Initialize MongoDB
try:
    client = pymongo.MongoClient(MONGO_URL, uuidRepresentation="standard")
    db = client["chat_with_doc"]
    conversation_collection = db["chat-history"]
    vector_collection = db["vector-store"]
    conversation_collection.create_index([("session_id")], unique=True)
except:
    print(traceback.format_exc())
    exc_type, exc_obj, exc_tb = sys.exc_info()
    fname = os.path.split(exc_tb.tb_frame.f_code.co_filename)[1]
    print(exc_type, fname, exc_tb.tb_lineno)


# Pydantic models
class ChatMessageSent(BaseModel):
    session_id: str = None
    user_input: str


# Helper functions
# def read_file_from_s3(file_name: str) -> str:
#     """Read and extract text content from S3 files"""
#     file_content = s3.get_object(Bucket=S3_BUCKET, Key=f"{S3_PATH}{file_name}")[
#         "Body"
#     ].read()

#     if file_name.endswith(".pdf"):
#         pdf_reader = PdfReader(BytesIO(file_content))
#         return " ".join(page.extract_text() for page in pdf_reader.pages)

#     elif file_name.endswith(".docx"):
#         docx_reader = docx.Document(BytesIO(file_content))
#         return "\n".join(paragraph.text for paragraph in docx_reader.paragraphs)

#     raise ValueError("Unsupported file format. Please use PDF or DOCX files.")


def ensure_search_index_exists(collection):
    """Check if index exists and create if it doesn't"""
    existing_indexes = collection.list_search_indexes()
    
    # Check if vector_index already exists
    index_exists = any(index['name'] == 'vector_index' for index in existing_indexes)
    
    if not index_exists:
        search_index_model = SearchIndexModel(
            definition={
                "fields": [
                    {
                        "type": "vector",
                        "path": "embedding",
                        "numDimensions": 1536,
                        "similarity": "cosine",
                    },
                    {"type": "filter", "path": "page"},
                ]
            },
            name="vector_index",
            type="vectorSearch",
        )
        collection.create_search_index(model=search_index_model)
        
def get_session() -> str:
    return str(uuid.uuid4())


def add_session_history(session_id: str, new_values: List):
    document = conversation_collection.find_one({"session_id": session_id})
    if document:
        conversation = document["conversation"]
        conversation.extend(new_values)
        conversation_collection.update_one(
            {"session_id": session_id}, {"$set": {"conversation": conversation}}
        )
    else:
        conversation_collection.insert_one(
            {"session_id": session_id, "conversation": new_values}
        )


def load_memory_to_pass(session_id: str):
    data = conversation_collection.find_one({"session_id": session_id})
    history = []
    if data:
        data = data["conversation"]
        for x in range(0, len(data), 2):
            history.extend([(data[x], data[x + 1])])
    return history
app = FastAPI()


@app.middleware("http")
async def add_process_time_header(request: Request, call_next):
    start_time = time.time()
    response = await call_next(request)
    process_time = time.time() - start_time
    # Log the timing
    print(f"Time taken: {process_time:.3f} seconds")
    # Add timing to response headers
    response.headers["X-Process-Time"] = str(process_time)
    return response


app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.post("/uploadFile")
async def uploadtos3(data_file: UploadFile):
    
    file_content = await data_file.read()
    print('passed the file read')
    file_id = str(uuid.uuid4())
    
    if data_file.filename.endswith(".pdf"):
        pdf_reader = PdfReader(BytesIO(file_content))
        text_content = " ".join(page.extract_text() for page in pdf_reader.pages)

    elif data_file.filename.endswith(".docx"):
        docx_reader = docx.Document(BytesIO(file_content))
        text_content = "\n".join(paragraph.text for paragraph in docx_reader.paragraphs)
    print('passed the file conversion')
    # Create new vectors
    data = [LangchainDocument(page_content=text_content, metadata={"file_id": file_id})]
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000, chunk_overlap=100, separators=["\n", " ", ""]
    )
    all_splits = text_splitter.split_documents(data)
    
    print('passed the text splitting')
    vector_store = MongoDBAtlasVectorSearch.from_documents(
        documents=all_splits,
        embedding=embeddings,
        index_name = "vector_index",
        collection=vector_collection,
    )
    print('c the vector store')
    ensure_search_index_exists(vector_collection)
    print('passed the search index')
    if vector_store:
        return {"file_path": file_id}
@app.post("/chat")
async def create_chat_message(chats: ChatMessageSent):
    try:
        session_id = chats.session_id or get_session()
        query = chats.user_input
        
        vector_store = MongoDBAtlasVectorSearch(
            collection=vector_collection,
            embedding=embeddings,
            index_name="vector_index"
        )

        retriever = vector_store.as_retriever(
            search_type="similarity_score_threshold",
            search_kwargs={"k": 5, "score_threshold": 0.05},
        )
        
        qa_chain = ConversationalRetrievalChain.from_llm(
            llm,
            retriever=retriever,
            condense_question_prompt=PromptTemplate.from_template(
                "Use the following pieces of context to answer the question about document at the end. "
                "Be direct and precise. If the information is not in the document, clearly state that. "
                "Chat History: {chat_history}\n"
                "Question: {question}"
            ),
            return_source_documents=True,
            verbose=True
        )
        
        with get_openai_callback() as cb:
            answer = qa_chain.invoke(
                {"question": query, "chat_history": load_memory_to_pass(session_id=session_id)}
            )
            token_usage = cb.total_tokens
            answer["total_tokens_used"] = token_usage
        return {"response": answer, "session_id": session_id}
    except Exception:
        print(traceback.format_exc())
        exc_type, exc_obj, exc_tb = sys.exc_info()
        fname = os.path.split(exc_tb.tb_frame.f_code.co_filename)[1]
        print(exc_type, fname, exc_tb.tb_lineno)
        raise HTTPException(status_code=status.HTTP_204_NO_CONTENT, detail="error")

import uvicorn


if __name__ == "__main__":
    uvicorn.run("backend:app", host="0.0.0.0", port=8000, reload=True)